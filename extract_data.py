import docx2txt
from docx import Document
import zipfile
import pandas as pd
import csv
import re
import os
import cv2
import numpy as np
import json

def extract_images(docx_path,output_folder, csv_path):
    doc = Document(docx_path)
    if not doc.inline_shapes:
        os.makedirs(output_folder, exist_ok=True)
        # This will extract images and return the plain text
        _ = docx2txt.process(docx_path, output_folder)

        images = [f for f in os.listdir(output_folder) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif'))]
        images.sort(key=lambda x: int(''.join(filter(str.isdigit, x))))  # Sort numerically if needed

        if not images:
            print("No images found.")
            return

        print(f"Extracted {len(images)} images in (almost) document order into '{output_folder}'.")

        # Rename logic
        custom_names = [
            "Swallow Composite",
            "Resting Pressure Profile & Anatomy",
            "Landmark Id"
        ]

        for i, image in enumerate(images):
            if i < len(custom_names):
                new_name = custom_names[i]
            else:
                new_name = f"Swallow #{i - len(custom_names) + 1}"

            # Clean up the filename
            safe_name = ''.join(c if c.isalnum() or c in [' ', '-', '_', '#'] else '_' for c in new_name).strip()
            new_filename = f"{safe_name}.png"
            
            old_path = os.path.join(output_folder, image)
            new_path = os.path.join(output_folder, new_filename)
            
            os.rename(old_path, new_path)

        return

    print(f"Found {len(doc.inline_shapes)} images/diagrams in the document. Extracting...")
    
    def get_image_names_from_csv(csv_path):
        df = pd.read_csv(csv_path)
        # Include headers (column names) and values
        headers = df.columns.tolist()
        values = df.values.flatten().tolist()
        all_names = headers + values
        all_names = [name.strip() for name in all_names if pd.notna(name)]
        return all_names
    
    csv_image_names = get_image_names_from_csv(csv_path)
    custom_image_names = [
        "Swallow Composite",
        "Resting Pressure Profile & Anatomy"
    ]

    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        
        images_folder = "images"
        os.makedirs(images_folder, exist_ok=True)
        
        rels = doc._part.rels
        for i, shape in enumerate(doc.inline_shapes):
            r_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = rels[r_id].target_part

            if i < 2:
                image_name = custom_image_names[i]
            else:
                index = i - 2
                image_name = csv_image_names[index] if index < len(csv_image_names) else f"Swallow_{i+1}"

            filename = os.path.join(images_folder, f"{image_name}.png")

            with open(filename, "wb") as f:
                f.write(image_part.blob)

    print(f"Saved {len(doc.inline_shapes)} images")


def process_docx_tables(docx_path, output_dir="."):
    """
    Extract and clean tables from a Word document, saving only the final processed CSV files.
    
    Args:
        docx_path (str): Path to the Word document
        output_dir (str): Directory to save output CSV files (default: current directory)
    """
    os.makedirs(output_dir, exist_ok=True)
    
    # Helper function to extract raw tables into memory
    def extract_raw_tables(file_path):
        doc = Document(file_path)
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        print(f"Found {len(tables_data)} tables from document")
        return tables_data

    # Clean table 1 (Patient data)
    def clean_table_1(table_data, output_path):
        # Convert table data to CSV-like string
        csv_string = "\n".join([",".join(row) for row in table_data])
        lines = [line.strip().replace('"', '') for line in csv_string.split("\n") if line.strip()]
        
        data = []
        patient_name = ""
        patient_id = ""
        
        for line in lines:
            if line.startswith("Patient"):
                continue
            if "AIGG" in line:
                parts = line.split(",")
                patient_id = parts[0].strip()
                rest = parts[1:]
                i = 0
                while i < len(rest):
                    key = rest[i].strip().replace(":", "").replace(";", "")
                    value = rest[i + 1].strip() if i + 1 < len(rest) else ""
                    data.append((key, value))
                    i += 2
            else:
                patient_name = line.strip()
        
        cleaned_data = [("Patient", patient_name), ("ID", patient_id)] + data
        df = pd.DataFrame(cleaned_data, columns=["Field", "Value"])
        df.to_csv(output_path, index=False)
        print(f"✅ Cleaned patient data into {output_path}")

    # Clean table 2 (Categorized data)
    def clean_table_2(table_data, output_path):
        # Convert table data to CSV-like string
        csv_string = "\n".join([" ".join(row) for row in table_data if any(row)]).strip()
        csv_string = re.sub(r"(?<!\n)([A-Z][A-Za-z ]+\*)", r"\n\1", csv_string)
        blocks = re.split(r"\n([A-Za-z ]+\*)\n", csv_string)
        
        cleaned_rows = []
        for i in range(1, len(blocks), 2):
            category = blocks[i].replace("*", "").strip()
            lines = blocks[i+1].strip().split("\n")
            for line in lines:
                if "\t" in line:
                    param, value = line.split("\t", 1)
                    cleaned_rows.append([category, param.strip(), value.strip()])
        
        with open(output_path, "w", newline="", encoding="utf-8") as outfile:
            writer = csv.writer(outfile)
            writer.writerow(["Category", "Parameter", "Value"])
            writer.writerows(cleaned_rows)
        print(f"✅ Cleaned {len(cleaned_rows)} rows into {output_path}")

    # Clean table 3 (Esophageal data)
    def clean_table_3(table_data):
        # Convert table data to DataFrame
        df = pd.DataFrame(table_data)
        
        # Lower Esophageal Sphincter
        lower_df = df.iloc[1:17, [0, 1, 2]].copy()
        lower_df.columns = ["Parameter", "Value", "Normal Range"]
        lower_path = f"{output_dir}/Lower_Esophageal_Sphincter.csv"
        lower_df.to_csv(lower_path, index=False)
        
        # Esophageal Motility
        motility_df = df.iloc[0:19, [3, 4, 5]].copy()
        motility_df.columns = ["Parameter", "Value", "Normal Range"]
        motility_path = f"{output_dir}/Esophageal_Motility.csv"
        motility_df.to_csv(motility_path, index=False)
        
        # Find split point for second table
        split_index = df[df.iloc[:, 0] == "Upper Esophageal Sphincter"].index[0]
        df2 = df.iloc[split_index:]
        df2.columns = ["Upper Esophageal Sphincter", "Value", "Normal",
                      "Pharyngeal / UES Motility", "Value_2", "Normal_2"]
        
        # Upper Esophageal Sphincter
        upper_df = df2.iloc[1:7, [0, 1, 2]].copy()
        upper_df.columns = ["Parameter", "Value", "Normal Range"]
        upper_path = f"{output_dir}/Upper_Esophageal_Sphincter.csv"
        upper_df.to_csv(upper_path, index=False)
        
        # Pharyngeal/UES Motility
        pharyngeal_df = df2.iloc[0:20, [3, 4, 5]].copy()
        pharyngeal_df.columns = ["Parameter", "Value", "Normal Range"]
        pharyngeal_path = f"{output_dir}/Pharyngeal_UES_Motility.csv"
        pharyngeal_df.to_csv(pharyngeal_path, index=False)
        
        print(f"✅ Split table 3 into: {lower_path}, {motility_path}, {upper_path}, {pharyngeal_path}")

    def save_last_table(table_data, output_path):
            if output_path:
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                with open(output_path, "w", newline="", encoding="utf-8") as csvfile:
                    writer = csv.writer(csvfile)
                    writer.writerows(table_data)
                print(f"✅ Saved last table to {output_path} for image names")

    # Execute the pipeline
    tables = extract_raw_tables(docx_path)
    
    # Process each table directly from memory
    if len(tables) >= 1:
        clean_table_1(tables[0], f"{output_dir}/Patient_details.csv")
    if len(tables) >= 2:
        clean_table_2(tables[1], f"{output_dir}/Esophageal_Manometry_Summary.csv")
    if len(tables) >= 3:
        clean_table_3(tables[2])
    save_last_table(tables[-1], f"{output_dir}/Image_filenames.csv")

def extract_text_to_json(file_path):
    sections = [
    "Chicago Classification Findings*",
    "Procedure",
    "Indications",
    "Interpretation / Findings",
    "Impressions"
]
    doc = Document(file_path)
    text = '\n'.join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

    extracted = {}
    for i, section in enumerate(sections):
        # Create a regex pattern to capture from current section to next section title
        start = re.escape(section)
        end = re.escape(sections[i+1]) if i + 1 < len(sections) else '$'
        pattern = rf"{start}\n(.*?)(?=\n{end})"
        match = re.search(pattern, text, re.DOTALL)
        if match:
            extracted[section] = match.group(1).strip()
    
    # Create the folder if it doesn't exist
    output_dir = "extracted_data"
    os.makedirs(output_dir, exist_ok=True)

    # Save JSON to file
    output_path = os.path.join(output_dir, "chicago_classification_findings.json")
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(extracted, f, indent=2, ensure_ascii=False)

    
    print(f"✅ Extracted text to {output_path} as JSON")

def remove_grid_lines_from_images(folder_path):
    # Get list of image files in the folder
    image_files = sorted([
        f for f in os.listdir(folder_path)
        if f.lower().endswith(('.png', '.jpg', '.jpeg'))
    ])

    # Skip the second image (index 1)
    image_files_to_process = [f for i, f in enumerate(image_files) if i != 1]

    # Create output folder in parent directory of input folder
    parent_folder = os.path.dirname(folder_path)
    output_folder = os.path.join(parent_folder, 'processed_images')
    os.makedirs(output_folder, exist_ok=True)

    for image_name in image_files_to_process:
        image_path = os.path.join(folder_path, image_name)
        img = cv2.imread(image_path)

        if img is None:
            print(f"Skipping invalid image: {image_name}")
            continue

        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        edges = cv2.Canny(gray, 50, 150, apertureSize=3)

        lines = cv2.HoughLinesP(edges, 1, np.pi / 180, threshold=30, minLineLength=5, maxLineGap=10)
        mask = np.zeros_like(gray)

        if lines is not None:
            for line in lines:
                x1, y1, x2, y2 = line[0]
                if abs(x1 - x2) < 5 or abs(y1 - y2) < 5:
                    cv2.line(mask, (x1, y1), (x2, y2), 255, 2)

        inpainted = cv2.inpaint(img, mask, inpaintRadius=3, flags=cv2.INPAINT_TELEA)

        # Save to processed_images folder in parent directory
        output_path = os.path.join(output_folder, image_name)
        cv2.imwrite(output_path, inpainted)
    print(f"Processed and saved all images at processed_images")

    
if __name__ == "__main__":
    docx_path = "subj2.docx"
    process_docx_tables(docx_path, "extracted_data")
    extract_text_to_json(docx_path)
    image_file_names = "extracted_data/Image_filenames.csv"    
    extract_images(docx_path, "images", "extracted_data/Image_filenames.csv" )
    remove_grid_lines_from_images('images')
